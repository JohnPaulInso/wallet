from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "SmartWallet_Project_Report.md"
OUTPUT = ROOT / "SmartWallet_Project_Report.docx"


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def add_subtitle(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def build_docx() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    add_title(document, "Smart Wallet Project Report")
    add_subtitle(document, "Detailed product and AI integration overview")
    document.add_paragraph()

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        document.add_paragraph(line)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_docx()
    print(path)
